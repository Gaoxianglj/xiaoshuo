from copy import deepcopy
from pathlib import Path

from docx import Document

from docx_fangsong import embed_fangsong_docx


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "正文" / "学院魔女_第三章_决斗与暗流.docx"
SOURCE = ROOT / "正文" / "正文MD" / "学院魔女_第四章_背叛.md"
OUTPUT = ROOT / "正文" / "学院魔女_第四章_背叛.docx"


def clone_paragraph_format(source, target):
    target.style = source.style
    target.alignment = source.alignment
    target_p = target.paragraph_format.element
    if target_p.pPr is not None:
        target_p.remove(target_p.pPr)
    if source.paragraph_format.element.pPr is not None:
        target_p.insert(0, deepcopy(source.paragraph_format.element.pPr))


def clone_run_format(source, target):
    if source._element.rPr is not None:
        target._element.insert(0, deepcopy(source._element.rPr))


def add_from_template(document, text, paragraph_template, run_template):
    paragraph = document.add_paragraph()
    clone_paragraph_format(paragraph_template, paragraph)
    run = paragraph.add_run(text)
    clone_run_format(run_template, run)


def main():
    document = Document(REFERENCE)
    title_template = document.paragraphs[0]
    section_template = document.paragraphs[1]
    body_template = document.paragraphs[2]
    title_run = title_template.runs[0]
    section_run = section_template.runs[0]
    body_run = body_template.runs[0]

    body = document._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            add_from_template(document, stripped[2:], title_template, title_run)
        elif stripped.startswith("## "):
            add_from_template(document, stripped[3:], section_template, section_run)
        else:
            add_from_template(document, stripped, body_template, body_run)

    document.save(OUTPUT)
    embed_fangsong_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
