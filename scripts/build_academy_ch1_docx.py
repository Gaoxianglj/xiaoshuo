from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docx_fangsong import FONT_NAME, embed_fangsong_docx


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "正文" / "正文MD" / "学院魔女_第一章_学院的新生.md"
OUTPUT = ROOT / "正文" / "学院魔女_第一章_学院的新生.docx"


def set_run_font(run, east_asia=FONT_NAME, latin=FONT_NAME, size=11, bold=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), latin)


def set_widow_control(paragraph, enabled=True):
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)
    widow.set(qn("w:val"), "1" if enabled else "0")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    set_run_font(run, size=9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])

    tail = paragraph.add_run(" —")
    set_run_font(tail, size=9)
    tail.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    add_page_number(section.footer.paragraphs[0])


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=18, bold=True)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=11.5, bold=True)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.333
    set_widow_control(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "第一章 学院的新生"
    doc.core_properties.subject = "《学院魔女》学院篇正文"
    doc.core_properties.author = ""

    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
        elif stripped.startswith("## "):
            add_section_heading(doc, stripped[3:].strip())
        else:
            add_body(doc, stripped)

    doc.save(OUTPUT)
    embed_fangsong_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
