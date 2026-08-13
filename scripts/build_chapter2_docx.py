from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docx_fangsong import FONT_NAME, embed_fangsong_docx


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "正文" / "学院魔女_第一章_学院的新生.docx"
MARKDOWN = ROOT / "正文" / "正文MD" / "学院魔女_第二章_唯一的友人.md"
OUTPUT = ROOT / "正文" / "学院魔女_第二章_唯一的友人.docx"


def set_font(run, size=12, bold=False):
    run.font.name = FONT_NAME
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for field in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{field}"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


doc = Document(SOURCE)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

normal = doc.styles["Normal"]
normal.font.name = FONT_NAME
for field in ("ascii", "hAnsi", "eastAsia", "cs"):
    normal._element.get_or_add_rPr().rFonts.set(qn(f"w:{field}"), FONT_NAME)
normal.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
for raw in lines:
    text = raw.strip()
    if not text:
        continue

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5

    if text.startswith("# "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        run = paragraph.add_run(text[2:])
        set_font(run, size=16, bold=True)
    elif text.startswith("## "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(12)
        run = paragraph.add_run(text[3:])
        set_font(run, size=12, bold=False)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        run = paragraph.add_run(text)
        set_font(run, size=12, bold=False)

doc.save(OUTPUT)
embed_fangsong_docx(OUTPUT)
print(OUTPUT)
